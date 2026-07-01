#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""最终插入 — 俞皓译分两格 + 模块图蓝底"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os, json

DOCX = r'D:\Visual Studio\Project\学生管理系统实验报告_最终版.docx'
OUT = r'D:\Visual Studio\Project\学生管理系统实验报告_完成.docx'
FLOW = r'D:\Visual Studio\Project\flowcharts_mermaid'

with open(os.path.join(FLOW, 'layout_v7.json')) as f:
    LAYOUT = json.load(f)
# 模块图尺寸从JSON读取，不覆盖

def find_cell(table, text):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if text in cell.text:
                return cell
    return None

def insert_img(cell, name, width, caption, page_break):
    path = os.path.join(FLOW, f'{name}.png')
    if not os.path.exists(path): return False
    if page_break:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.add_run().add_break(WD_BREAK.PAGE)
    sp = cell.add_paragraph()
    sp.paragraph_format.space_before = Pt(8)
    ip = cell.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(path, width=Inches(width))
    cp = cell.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(caption).font.size = Pt(9)
    return True

def main():
    doc = Document(DOCX)
    t = doc.tables[1]

    plans = [
        # (keyword, image, person, use_row10_for_extra)
        ('程序整体的模块结构图', '01_模块结构图', '公共', None),
        ('【模块一：主菜单循环与文件I/O', '02_主程序流程图', '宫士淇', None),
        ('【模块二：排序统计与导出（SortAndExport）】', '03_SortAndExport流程图', '宫士淇', None),
        ('【模块一：添加学生记录 Insert()】', '04_Insert流程图', '王逸轩', None),
        ('【模块二：浏览全部学生 TraverseTable()】', '05_TraverseTable流程图', '王逸轩', None),
        # 俞皓译分两格: 前2张在Row9(设计), 后2张在Row10(总结)
        ('【模块一：按学号删除 Delete()】', '06_Delete流程图', '俞皓译', None),
        ('SearchById 算法（散列定位，高效）：', '07_SearchById流程图', '俞皓译', None),
        ('SearchByName 算法（全表扫描）：', '08_SearchByName流程图', '俞皓译', '实践总结（俞皓译'),
        ('【模块三：按学号修改 Update()】', '09_Update流程图', '俞皓译', '实践总结（俞皓译'),
    ]

    extra_after = {}

    ok = 0
    for kw, name, person, alt_kw in plans:
        if name not in LAYOUT: continue
        w, h, pb = LAYOUT[name]

        # 如果有备用关键字, 用备用位置
        if alt_kw:
            cell = find_cell(t, alt_kw)
            if cell:
                print(f'  {name}: 插入到实践总结格 [{person}]')
            else:
                cell = find_cell(t, kw)
        else:
            cell = find_cell(t, kw)

        if not cell:
            print(f'  [未找到] {name}')
            continue

        insert_img(cell, name, w, f'图 {name}', pb)
        print(f'  {name}: {w:.1f}\"x{h:.1f}\" [{person}]', end='')
        if pb: print(' +分页', end='')
        print()
        ok += 1

        # 紧随后续图
        if name in extra_after:
            en = extra_after[name]
            if en in LAYOUT:
                ew, eh, epb = LAYOUT[en]
                insert_img(cell, en, ew, f'图 {en}', epb)
                print(f'    +{en}: {ew:.1f}\"x{eh:.1f}\"')
                ok += 1

    doc.save(OUT)
    print(f'\n{ok}张 -> {OUT} ({os.path.getsize(OUT)//1024}KB)')

if __name__ == '__main__':
    main()
