"""
Update 学生管理系统实验报告_完成版.docx with:
1. Single-person completion (独立完成)
2. Reorganized code (main at top)
3. Latest test results
"""
from docx import Document
import re

doc = Document('程序设计实践报告模板.docx')

# ============================================================
# 1. Update Cover Page
# ============================================================
for p in doc.paragraphs:
    if '课题名称' in p.text:
        for run in p.runs:
            if '课题名称' in run.text:
                run.text = '课题名称：学生信息管理系统（基于散列表）'
                break
    elif '班    级' in p.text:
        for run in p.runs:
            if '班    级' in run.text:
                run.text = '班    级   计科（国教）B25—1    '
                break
    elif '完成日期' in p.text:
        for run in p.runs:
            if '完成日期' in run.text:
                run.text = '完成日期：  2026年6月10日        '
                break

# ============================================================
# 2. Update Table 0 - Single person (独立完成)
# ============================================================
table0 = doc.tables[0]
# Row 1: 组长 -> 独立完成
table0.cell(1, 0).text = '独立完成'
table0.cell(1, 1).text = '2520109'
table0.cell(1, 2).text = '宫士淇'
table0.cell(1, 3).text = '100'
# Rows 2-3: Clear the "组员" rows
for r in [2, 3]:
    for c in range(4):
        table0.cell(r, c).text = ''

# ============================================================
# 3. Update Table 1 Row 0 - Main content cell
# ============================================================
table1 = doc.tables[1]
main_cell = table1.cell(0, 0)
main_text = main_cell.text

# --- Replace storage structure in 三、总体设计 ---
old_storage = r'​?本系统采用顺序存储结构（结构体数组）来保存学生数据。\n// 定义学生信息结构体\n.*?St\.list'
new_storage = """本系统采用散列表（哈希表）的链地址法来存储学生数据。散列函数采用除留余数法：HashFunc(id) = id % TABLE_SIZE，其中 TABLE_SIZE = 13。当发生散列冲突时，同一散列地址的学生记录以单链表的形式链接在一起（头插法）。

// 函数结果状态代码
#define   OK    1
#define   ERROR 0
#define   OVERFLOW -2
typedef int Status;

#define TABLE_SIZE 13

// 学生记录结构体
typedef struct {
    int id;           // 学号：唯一标识
    char name[20];    // 姓名
    int gender;       // 性别，0为男生，1为女生
    int age;          // 年龄
    float score;      // 成绩
} Student;

// 散列表结点（链地址法）
typedef struct HashNode {
    Student data;           // 学生数据
    struct HashNode* next;  // 指向下一个同义词结点
} HashNode;

// 散列表：指针数组，每个元素指向一个链表（桶）
typedef HashNode* HashTable[TABLE_SIZE];"""

main_text = re.sub(old_storage, new_storage, main_text, flags=re.DOTALL)

# --- Replace function list in 三、总体设计 ---
old_func = """3、所实现的功能函数
（每个功能对应的函数名、函数形参及返回值说明）"""

new_func = """3、所实现的功能函数
（每个功能对应的函数名、函数形参及返回值说明）

本系统共实现以下12个功能函数：

（1）int HashFunc(int id)
    功能：散列函数（除留余数法），计算学号对应的散列地址
    参数：id — 学生学号
    返回值：散列地址（0 ~ TABLE_SIZE-1）

（2）Status InitTable(HashTable& H)
    功能：初始化散列表，将所有桶置为空指针
    参数：H — 散列表（引用传递）
    返回值：OK

（3）Student* SearchById(HashTable H, int id)
    功能：按学号精确查找学生记录，利用散列函数直接定位到桶
    参数：H — 散列表，id — 待查找的学号
    返回值：成功返回指向学生数据的指针，失败返回NULL

（4）Student* SearchByName(HashTable H, const char* name)
    功能：按姓名查找学生记录（全表扫描，返回第一个匹配项）
    参数：H — 散列表，name — 待查找的姓名
    返回值：成功返回指向学生数据的指针，失败返回NULL

（5）Status Insert(HashTable& H, Student stu)
    功能：插入一条学生记录，含学号查重，采用头插法
    参数：H — 散列表（引用），stu — 待插入的学生数据
    返回值：OK（成功）/ ERROR（学号重复）/ OVERFLOW（内存不足）

（6）Status Delete(HashTable& H, int id)
    功能：根据学号删除指定学生的整条记录
    参数：H — 散列表（引用），id — 待删除的学号
    返回值：OK（成功）/ ERROR（未找到）

（7）Status Update(HashTable& H, int id, Student newStu)
    功能：根据学号定位并修改学生信息，支持修改学号（先删后插）
    参数：H — 散列表（引用），id — 原学号，newStu — 新数据
    返回值：OK（成功）/ ERROR（未找到或新学号冲突）

（8）void TraverseTable(HashTable H)
    功能：以规范的表格形式（含分隔线和列标题）输出所有学生
    参数：H — 散列表

（9）Status LoadFromFile(HashTable& H, const char* filename)
    功能：从数据文件中逐行读取学生记录并插入散列表
    参数：H — 散列表（引用），filename — 数据文件名
    返回值：OK（成功）/ ERROR（文件不存在）

（10）Status SaveToFile(HashTable H, const char* filename)
    功能：将散列表中所有学生记录保存到数据文件
    参数：H — 散列表，filename — 目标文件名
    返回值：OK（成功）/ ERROR（无法写入）

（11）int CountStudents(HashTable H)
    功能：遍历散列表统计当前学生记录总数
    参数：H — 散列表
    返回值：学生总数

（12）Status SortAndExport(HashTable H, const char* filename)
    功能：提取→排序（性别升序+成绩降序）→统计→导出到文件
    参数：H — 散列表，filename — 导出文件名
    返回值：OK（成功）/ ERROR（无记录）"""

main_text = main_text.replace(old_func, new_func)

# --- Replace test section in 五、测试 ---
old_test = """五、 测试
（整合各功能模块后的测试结果截图及说明）"""

new_test = """五、 测试
（整合各功能模块后的测试结果截图及说明）

以下为程序实际运行后各功能模块的测试结果。

【测试1：程序启动与数据加载】
程序启动时自动从 records.txt 加载8条学生记录，控制台输出：

已从 records.txt 加载 8 条学生记录。

主菜单界面显示如下：
========================================
       学生信息管理系统（散列表）
========================================
  1. 添加学生记录
  2. 删除学生记录
  3. 查找学生记录
  4. 修改学生记录
  5. 浏览全部学生
  6. 排序统计与导出
  0. 退出系统
========================================

【测试2：数据浏览功能（菜单选项5）】
以表格形式完整输出当前8条记录：
---------------------------------------------------------
学号    姓名      性别  年龄  成绩
---------------------------------------------------------
1         Wu          女     19      64.0
2         Zhao        女     20      58.0
4         Qian        男     21      81.0
5         Li          男     20      68.0
6         Wang        女     20      72.0
7         Zheng       女     18      55.0
8         Sun         男     19      74.0
9         Zhou        男     22      75.0
---------------------------------------------------------

【测试3：双向查找功能（菜单选项3）】
（a）按学号查找（学号=1）：
---------- 查找结果 ----------
学号：1
姓名：Wu
性别：女
年龄：19
成绩：64.0
------------------------------

（b）按姓名查找（姓名=Wu）：
查找结果同上，双向查询均能正确定位。

【测试4：添加学生与学号查重（菜单选项1）】
（a）添加新学生（学号=10，姓名=ZhangSan，性别=男，年龄=22，成绩=88.5）：
成功添加学生 ZhangSan（学号：10）

添加后浏览表格，共9条记录，新增记录已正确显示在第9行。

（b）尝试重复添加已存在的学号（学号=1）：
错误：学号 1 已存在，插入失败！
学号查重功能验证通过。

【测试5：删除学生记录（菜单选项2）】
删除学号为10的学生（ZhangSan）：
成功删除学生 ZhangSan（学号：10）

删除后浏览表格，已恢复为8条记录，ZhangSan的记录已完全移除。

【测试6：排序统计与导出（菜单选项6）】
排序验证——男生在前（按成绩降序）：
  Qian 81.0, Zhou 75.0, Sun 74.0, Li 68.0
排序验证——女生在后（按成绩降序）：
  Wang 72.0, Wu 64.0, Zhao 58.0, Zheng 55.0

统计结果输出：
========== 统计结果 ==========
男生人数：4，合格人数：4，合格率：100.0%
女生人数：4，合格人数：2，合格率：50.0%
全班平均分：68.38
==============================

排序后的完整报表已保存至 sorted_output.txt。导出的文件内容包含排序后的学生记录表格及上述三项统计指标，格式与数据均正确无误。"""

main_text = main_text.replace(old_test, new_test)

# Write back to main cell
main_cell.text = main_text

# ============================================================
# 4. Fill in the personal sections - All by 宫士淇 (独立完成)
# ============================================================

# --- Section labels (学号/姓名 rows) ---
table1.cell(2, 0).text = '2520109'
table1.cell(2, 1).text = '2520109'
table1.cell(2, 2).text = '2520109'
table1.cell(2, 3).text = '2520109'
table1.cell(2, 7).text = '宫士淇'
table1.cell(2, 8).text = '宫士淇'
table1.cell(2, 9).text = '宫士淇'

# Clear rows 5-10 (old 组员1 and 组员2 sections)
for r in range(5, 11):
    for c in range(12):
        try:
            table1.cell(r, c).text = ''
        except:
            pass

# --- 宫士淇: Module 1 - 系统主框架与排序统计导出 (Row 3, content cell Col 2) ---
design_m1 = """1、算法设计 —— 系统主框架、文件I/O与排序统计导出模块

【模块一：主菜单循环与文件I/O（main、LoadFromFile、SaveToFile）】
算法流程：
(1) 程序启动 → 调用 InitTable() 将散列表所有桶置空
(2) 调用 LoadFromFile("records.txt") 从数据文件加载学生记录：
    a. 打开文件，若不存在则提示并返回ERROR
    b. 循环读取每行数据：id → name → gender → age → score
    c. 每读取一条即调用 Insert() 插入散列表（自带查重）
    d. 关闭文件，输出加载记录数
(3) 进入主循环 while(true)：
    a. 调用 showMenu() 显示7选项菜单
    b. cin >> choice 读取用户选择
    c. 检测 cin.eof() || cin.fail() 防止管道输入时死循环
    d. 调用 clearInput() 清空输入缓冲区
    e. switch-case 分发：1→addStudent, 2→deleteStudent, 3→searchMenu,
       4→updateStudent, 5→TraverseTable, 6→SortAndExport, 0→退出
(4) 选择0时调用 SaveToFile("records.txt") 保存所有记录后退出

【模块二：排序统计与导出（SortAndExport）】
算法流程：
(1) 调用 CountStudents() 获取记录总数 total，若为0则返回ERROR
(2) 动态分配 Student 数组 arr[total]
(3) 双重循环遍历散列表所有桶的所有结点，提取数据到 arr[]
(4) 冒泡排序（双重循环）：
    外层 i: 0 → total-2，内层 j: 0 → total-2-i
    比较条件：
    - 若 arr[j].gender > arr[j+1].gender → 交换（性别升序，男0在前女1在后）
    - 若 gender 相同且 arr[j].score < arr[j+1].score → 交换（成绩降序）
(5) 遍历已排序数组，累计统计：
    sumScore（总分）、maleCount/malePass（男生人数/合格人数）、
    femaleCount/femalePass（女生人数/合格人数），合格标准为 score >= 60
(6) 计算三项统计指标：
    avgScore = sumScore / total
    maleRate = malePass / maleCount × 100（maleCount==0 则为 0）
    femaleRate = femalePass / femaleCount × 100
(7) 打开输出文件，写入排序后的学生表格（含表头和分隔线）及三项统计结果
(8) 同时将统计结果输出到控制台
(9) 释放动态数组 delete[] arr，返回 OK

2、代码设计

// 主函数 —— 菜单驱动的交互界面（main在顶部，函数定义在下方）
int main() {
    HashTable H;
    InitTable(H);
    // 启动时自动加载数据
    LoadFromFile(H, "records.txt");

    int choice;
    while (true) {
        showMenu();
        cin >> choice;
        // 检测EOF防止死循环
        if (cin.eof() || cin.fail()) {
            cout << endl << "输入结束，系统退出。" << endl;
            break;
        }
        clearInput();
        switch (choice) {
            case 1: addStudent(H); break;
            case 2: deleteStudent(H); break;
            case 3: searchMenu(H); break;
            case 4: updateStudent(H); break;
            case 5: TraverseTable(H); break;
            case 6: SortAndExport(H, "sorted_output.txt"); break;
            case 0:
                SaveToFile(H, "records.txt");
                cout << "数据已保存，感谢使用！" << endl;
                return 0;
            default:
                cout << "无效选项，请重新输入（0-6）！" << endl;
        }
    }
    return 0;
}

// 从文件加载数据
Status LoadFromFile(HashTable& H, const char* filename) {
    ifstream infile(filename);
    if (!infile.is_open()) {
        cout << "提示：数据文件 " << filename << " 不存在，将从空表开始。" << endl;
        return ERROR;
    }
    Student stu;
    int count = 0;
    while (infile >> stu.id >> stu.name >> stu.gender >> stu.age >> stu.score) {
        Insert(H, stu);
        count++;
    }
    infile.close();
    cout << "已从 " << filename << " 加载 " << count << " 条学生记录。" << endl;
    return OK;
}

// 排序、统计并导出到文件
Status SortAndExport(HashTable H, const char* filename) {
    int total = CountStudents(H);
    if (total == 0) {
        cout << "当前无学生记录，无法导出！" << endl;
        return ERROR;
    }
    // 1. 提取所有学生到动态数组
    Student* arr = new Student[total];
    int idx = 0;
    for (int i = 0; i < TABLE_SIZE; i++) {
        HashNode* p = H[i];
        while (p != NULL) {
            arr[idx++] = p->data;
            p = p->next;
        }
    }
    // 2. 冒泡排序：性别升序，同性别成绩降序
    for (int i = 0; i < total - 1; i++) {
        for (int j = 0; j < total - 1 - i; j++) {
            bool needSwap = false;
            if (arr[j].gender > arr[j + 1].gender) {
                needSwap = true;
            } else if (arr[j].gender == arr[j + 1].gender
                       && arr[j].score < arr[j + 1].score) {
                needSwap = true;
            }
            if (needSwap) {
                Student temp = arr[j];
                arr[j] = arr[j + 1];
                arr[j + 1] = temp;
            }
        }
    }
    // 3. 统计数据
    int maleCount = 0, malePass = 0;
    int femaleCount = 0, femalePass = 0;
    float sumScore = 0.0;
    for (int i = 0; i < total; i++) {
        sumScore += arr[i].score;
        if (arr[i].gender == 0) {
            maleCount++;
            if (arr[i].score >= 60.0) malePass++;
        } else {
            femaleCount++;
            if (arr[i].score >= 60.0) femalePass++;
        }
    }
    float avgScore = sumScore / total;
    float maleRate = (maleCount > 0) ? (float)malePass / maleCount * 100 : 0;
    float femaleRate = (femaleCount > 0) ? (float)femalePass / femaleCount * 100 : 0;
    // 4. 写入文件并输出到控制台
    ofstream outfile(filename);
    outfile << "==========================" << endl;
    outfile << "   排序后学生信息报表" << endl;
    outfile << "==========================" << endl;
    // ... 表格头和数据 ...
    outfile << "男生合格率（>=60分）：" << maleRate << "%" << endl;
    outfile << "女生合格率（>=60分）：" << femaleRate << "%" << endl;
    outfile << "全班平均分：" << avgScore << endl;
    outfile.close();
    delete[] arr;
    cout << "========== 统计结果 ==========" << endl;
    cout << "男生人数：" << maleCount << "，合格人数：" << malePass
         << "，合格率：" << maleRate << "%" << endl;
    cout << "女生人数：" << femaleCount << "，合格人数：" << femalePass
         << "，合格率：" << femaleRate << "%" << endl;
    cout << "全班平均分：" << avgScore << endl;
    return OK;
}"""

# --- 宫士淇: Module 2 - 基础数据管理 (Row 3, cell starting Col 4 or equivalent) ---
# Since cells are merged, we'll put Module 2 in a separate merged section
design_m2 = """1、算法设计 —— 基础数据管理模块（Insert、TraverseTable）

【模块一：添加学生记录 Insert()】
算法流程：
(1) 接收 Student 结构体参数 stu（含学号、姓名、性别、年龄、成绩）
(2) 学号查重：调用 SearchById(H, stu.id) 在对应散列桶的链表中查找
(3) 若 SearchById 返回非NULL（学号已存在），输出"错误：学号 XX 已存在，插入失败！"，返回 ERROR
(4) 若学号不重复，计算散列地址 addr = HashFunc(stu.id)
(5) 使用 new 动态分配 HashNode 结点，若分配失败（返回NULL）则返回 OVERFLOW
(6) 将 stu 数据复制到新结点的 data 域
(7) 头插法插入：p->next = H[addr]; H[addr] = p（时间复杂度 O(1)）
(8) 输出成功提示信息，返回 OK

【模块二：浏览全部学生 TraverseTable()】
算法流程：
(1) 输出表头分隔线（57个减号）
(2) 使用 iomanip 操纵符输出列标题行：setw(10)学号、setw(12)姓名、setw(8)性别、setw(8)年龄、setw(10)成绩
(3) 再次输出分隔线
(4) 设置空表标志 empty = true
(5) 外循环遍历 TABLE_SIZE=13 个桶 (i: 0→12)
(6) 内循环遍历每个桶的链表（p = H[i]; p != NULL; p = p->next）
(7) 对每个结点输出一行格式化数据：
    - 学号/姓名/年龄用 setw 固定列宽
    - 性别用三目运算符转换：gender==0 ? "男" : "女"
    - 成绩用 fixed + setprecision(1) 显示一位小数
    - 设置 empty = false
(8) 若 empty 仍为 true（所有桶均无结点），输出"（当前无学生记录）"
(9) 输出表尾分隔线

2、代码设计

// 插入学生记录，含学号查重
Status Insert(HashTable& H, Student stu) {
    // 学号查重——必须先检查再插入
    if (SearchById(H, stu.id) != NULL) {
        cout << "错误：学号 " << stu.id << " 已存在，插入失败！" << endl;
        return ERROR;
    }
    int addr = HashFunc(stu.id);
    HashNode* p = new HashNode;
    if (p == NULL) return OVERFLOW;  // 内存分配失败
    p->data = stu;
    // 头插法：新结点直接插入桶首，O(1)时间
    p->next = H[addr];
    H[addr] = p;
    cout << "成功添加学生 " << stu.name << "（学号：" << stu.id << "）" << endl;
    return OK;
}

// 以规范表格形式输出所有学生信息
void TraverseTable(HashTable H) {
    cout << endl;
    cout << "---------------------------------------------------------" << endl;
    cout << left << setw(10) << "学号"
         << setw(12) << "姓名"
         << setw(8) << "性别"
         << setw(8) << "年龄"
         << setw(10) << "成绩" << endl;
    cout << "---------------------------------------------------------" << endl;
    bool empty = true;
    for (int i = 0; i < TABLE_SIZE; i++) {
        HashNode* p = H[i];
        while (p != NULL) {
            empty = false;
            cout << left << setw(10) << p->data.id
                 << setw(12) << p->data.name
                 << setw(8) << (p->data.gender == 0 ? "男" : "女")
                 << setw(8) << p->data.age
                 << setw(10) << fixed << setprecision(1) << p->data.score << endl;
            p = p->next;
        }
    }
    if (empty) {
        cout << "（当前无学生记录）" << endl;
    }
    cout << "---------------------------------------------------------" << endl;
}"""

# --- 宫士淇: Module 3 - 数据维护与检索 (filling the last major cell) ---
design_m3 = """1、算法设计 —— 数据维护与检索模块（Delete、SearchById、SearchByName、Update）

【模块一：按学号删除 Delete()】
算法流程：
(1) 计算散列地址 addr = HashFunc(id)，直接定位目标桶
(2) 初始化两个指针：p = H[addr]（当前结点），prev = NULL（前驱结点）
(3) 在链表中遍历 while(p != NULL)：
    a. 若 p->data.id == id（找到目标）：
       - 若 prev == NULL（桶首结点）：H[addr] = p->next（更新桶指针）
       - 否则（链表中间/末尾）：prev->next = p->next（跳过被删结点）
       - 输出被删学生姓名和学号
       - delete p 释放内存
       - return OK
    b. 未找到则 prev = p; p = p->next 继续向后遍历
(4) 遍历完毕未找到：输出错误提示，返回 ERROR

【模块二：双向查询 SearchById / SearchByName】
SearchById 算法（散列定位，高效）：
(1) 计算 addr = HashFunc(id)，O(1) 定位到桶
(2) 仅遍历该桶的短链表，逐一比对 p->data.id
(3) 找到则返回 &(p->data)，未找到返回 NULL
(4) 平均时间复杂度 O(1)，最坏 O(n)（所有记录散列到同一桶）

SearchByName 算法（全表扫描）：
(1) 姓名非散列关键字，需遍历全部 TABLE_SIZE 个桶
(2) 外循环遍历桶，内循环用 strcmp 比对 p->data.name
(3) 返回第一个匹配项的指针，未找到返回 NULL
(4) 时间复杂度 O(n)

【模块三：按学号修改 Update()】
算法流程：
(1) 调用 SearchById(H, id) 定位目标学生，未找到则返回ERROR
(2) 记录找到的 Student* 指针 p
(3) 若 newStu.id != id（用户修改了学号）：
    a. 调用 SearchById(H, newStu.id) 检查新学号是否已存在
    b. 若新学号冲突 → 输出错误提示，返回 ERROR
    c. 若不冲突 → 调用 Delete(H, id) 删除旧记录
    d. 调用 Insert(H, newStu) 以新学号插入（散列地址可能变化）
(4) 若 newStu.id == id（学号不变）：
    a. 通过指针 p 直接原地修改：strcpy 更新姓名，赋值更新 gender/age/score
    b. 输出成功提示，返回 OK

2、代码设计

// 根据学号删除学生记录（带头结点前驱追踪）
Status Delete(HashTable& H, int id) {
    int addr = HashFunc(id);
    HashNode* p = H[addr];
    HashNode* prev = NULL;
    while (p != NULL) {
        if (p->data.id == id) {
            if (prev == NULL)
                H[addr] = p->next;   // 删除桶首结点
            else
                prev->next = p->next; // 删除中间/末尾结点
            cout << "成功删除学生 " << p->data.name
                 << "（学号：" << id << "）" << endl;
            delete p;                 // 释放内存
            return OK;
        }
        prev = p;
        p = p->next;
    }
    cout << "错误：未找到学号为 " << id << " 的学生！" << endl;
    return ERROR;
}

// 按学号查找（O(1)散列定位 + 短链表遍历）
Student* SearchById(HashTable H, int id) {
    int addr = HashFunc(id);
    HashNode* p = H[addr];
    while (p != NULL) {
        if (p->data.id == id)
            return &(p->data);
        p = p->next;
    }
    return NULL;
}

// 按姓名查找（全表扫描，返回第一个匹配项）
Student* SearchByName(HashTable H, const char* name) {
    for (int i = 0; i < TABLE_SIZE; i++) {
        HashNode* p = H[i];
        while (p != NULL) {
            if (strcmp(p->data.name, name) == 0)
                return &(p->data);
            p = p->next;
        }
    }
    return NULL;
}

// 根据学号修改学生信息（支持学号变更）
Status Update(HashTable& H, int id, Student newStu) {
    Student* p = SearchById(H, id);
    if (p == NULL) {
        cout << "错误：未找到学号为 " << id << " 的学生！" << endl;
        return ERROR;
    }
    // 若修改了学号，检查新学号是否重复
    if (newStu.id != id && SearchById(H, newStu.id) != NULL) {
        cout << "错误：新学号 " << newStu.id << " 已存在！" << endl;
        return ERROR;
    }
    // 学号改变→先删后插（散列地址可能变化）
    if (newStu.id != id) {
        Delete(H, id);
        Insert(H, newStu);
    } else {
        // 学号不变→原地修改，效率高
        strcpy(p->name, newStu.name);
        p->gender = newStu.gender;
        p->age = newStu.age;
        p->score = newStu.score;
        cout << "成功修改学生信息（学号：" << id << "）" << endl;
    }
    return OK;
}"""

# --- Unified practice summary ---
summary_all = """实践总结（宫士淇——独立完成全部模块）：

在本课题中，我独立完成了学生信息管理系统（基于散列表）的全部设计与开发工作，涵盖数据层（hash_table.h，共337行）和界面层（student_management.cpp，约289行）两大模块，实现了12个完整的功能函数。

一、系统架构设计

我选择了散列表的链地址法作为核心存储结构——以除留余数法 HashFunc(id)=id%13 计算散列地址，每个桶用单链表解决冲突。这种"数据层+界面层"的分离设计使代码结构清晰、易于维护。所有数据结构和操作函数封装在头文件 hash_table.h 中，主程序 student_management.cpp 负责用户交互界面调度，main函数放在文件顶部便于整体把握程序逻辑。

二、各模块实现要点

主框架模块：实现了清晰的7选项控制台菜单界面和主循环调度。为提升用户体验和健壮性，加入了输入验证函数（readInt、readFloat 带类型检查与错误恢复）和 clearInput 缓冲区清理函数。关键设计是管道测试时的EOF检测——通过 cin.eof() || cin.fail() 检测防止死循环。

文件I/O模块：实现了程序启动时自动加载和退出时自动保存的双向持久化机制，8条样例数据存储在 records.txt 中，格式为：学号 姓名 性别 年龄 成绩（空格分隔）。

Insert函数：核心在于学号唯一性保障——插入前先调用 SearchById 遍历对应散列桶的链表查重，查重通过后采用头插法将新结点插入桶首，时间复杂度为 O(1)。

Delete函数：通过维护 prev 前驱指针统一处理桶首删除和链表中间/末尾删除两种情况，删除后调用 delete 释放堆内存防止内存泄漏。

SearchById：充分利用散列表的核心优势——通过 HashFunc(id) 将查找范围从全表缩小到单个桶的短链表，平均查找长度远小于全表扫描。

Update函数：难点在于学号可变性处理——若用户修改了学号导致散列地址变化，采用"先Delete旧记录再Insert新记录"的策略，同时检查新学号是否已存在。

SortAndExport函数：是本课题的核心难点。由于散列表本身无序，设计了"提取→排序→统计→导出"四步走策略：先用动态数组收集所有学生记录，再用冒泡排序按性别升序（男0在前女1在后）+成绩降序的复合条件排序，接着遍历统计男女合格率（≥60分为合格）和全班平均分，最后将结果同时输出到控制台和文件。

三、调试中遇到的主要技术问题及解决措施

（1）散列函数命名冲突：最初命名为 hash，与 C++ 标准库 std::hash 模板类冲突导致编译失败（g++ 报 "reference to 'hash' is ambiguous"），解决方案是将函数改名为 HashFunc。

（2）管道测试死循环：echo "5" | ./program 导致无限循环，原因是 cin 到达 EOF 后 failbit 被置位但未检测，通过增加 cin.eof() || cin.fail() 检测并在检测到EOF时 break 解决。

（3）Update函数学号变化处理：若用户修改学号导致散列地址变化，简单的原地修改会破坏散列表结构，需要采用"先Delete旧记录再Insert新记录"的策略。

（4）TraverseTable表格对齐问题：初期未使用 setw 固定列宽导致表格错位，通过统一使用 setw 设定各列最小宽度并配合 left 左对齐解决。

（5）searchMenu缓冲区残留问题：cin >> choice 后缓冲区残留换行符会导致后续 getline 读到空字符串，通过在每个 >> 操作后统一调用 clearInput() 解决。

四、收获与不足

通过本次课程设计，我深入理解了散列表的存储原理与链地址法实现，掌握了模块化程序设计的完整流程（需求分析→数据结构设计→函数接口定义→编码实现→测试验证→文档撰写），并独立完成了全部代码与实验报告的撰写。

不足之处：（1）排序算法采用了O(n²)的冒泡排序，在数据量极大时效率偏低，后续可优化为快速排序或使用STL的sort；（2）散列表大小固定为13，未实现动态扩容（rehashing）机制；（3）按姓名查找属于顺序查找 O(n)，若需频繁使用可建立姓名到学号的辅助散列索引。"""

# Write content to cells - using the existing merged cell structure
# Module 1: system framework + sort/export -> Row 3, Col 2
table1.cell(3, 2).text = design_m1
# Summary -> Row 4, Col 2
table1.cell(4, 2).text = summary_all

# Module 2: Insert + TraverseTable -> Row 6, Col 2
table1.cell(6, 2).text = design_m2

# Module 3: Delete + Search + Update -> Row 9, Col 1
table1.cell(9, 1).text = design_m3

# ============================================================
# 5. Save
# ============================================================
output_path = '学生管理系统实验报告_完成版.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print('Done! All sections updated.')
