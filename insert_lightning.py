#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""将6张图插入避雷针调查报告Word文档（按顺序编号）并清理LaTeX"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os

SRC = r'D:\Visual Studio\Project\避雷针的静电防护原理调查报告.docx'
OUT = r'D:\Visual Studio\Project\避雷针的静电防护原理调查报告_完成.docx'

# 按文档出现顺序：图1~图6
IMAGES = [
    # 二.1 静电感应 → 图1
    ('大地内异种电荷会被吸引至地表位置',
     r'D:\Visual Studio\Project\diagrams\fig3_静电感应过程.png',
     5.0, '图1  静电感应与尖端放电过程'),
    # 二.1 尖端放电 → 图2
    ('产生放电现象',
     r'D:\Visual Studio\Project\diagrams\fig2_尖端放电机理.png',
     5.0, '图2  尖端放电机理：尖端 vs 钝面对比'),
    # 二.2 结构描述 → 图3
    ('避雷针由接闪尖端、接地引下线、接地装置三部分构成',
     r'D:\Visual Studio\Project\diagrams\fig1_避雷针结构图.png',
     4.0, '图3  避雷针系统结构与工作原理'),
    # 二.2 工作过程 → 图4
    ('当带有负电荷的雷雨云出现在建筑上方',
     r'D:\Visual Studio\Project\flowcharts_mermaid\05_避雷针工作流程图.png',
     3.2, '图4  避雷针工作全流程'),
    # 二.2 公式 → 图5
    ('设电场强度公式',
     r'D:\Visual Studio\Project\diagrams\fig6_公式汇总.png',
     4.5, '图5  避雷针静电防护核心公式与计算示例'),
    # 三 保护范围 → 图6
    ('避雷针存在固定保护区域',
     r'D:\Visual Studio\Project\diagrams\fig4_滚球法模型.png',
     4.5, '图6  滚球法保护范围几何模型'),
]


def add_citations(doc):
    """在正文对应位置标注引用序号"""
    # (关键词, 引用标记, 追加方式: 'end'=段末 'after'=关键词后)
    citations = [
        # 二.1 基础物理概念
        ('大地内异种电荷会被吸引至地表位置', '[1]', 'end'),       # 静电感应
        ('导体表面电荷面密度和曲率半径成反比', '[1]', 'end'),     # 尖端放电 σ∝1/r
        ('处于电场中的导体，内部电场强度最终为零', '[1]', 'end'), # 静电平衡
        # 二.2 工作模型
        ('避雷针由接闪尖端、接地引下线、接地装置三部分构成', '[4]', 'end'),  # 防雷规范结构定义
        ('电场强度极大，击穿周围空气，空气被电离形成导电通道', '[1]', 'end'), # 空气击穿原理
        ('设电场强度公式', '[1]', 'after'),                      # E=σ/ε₀ 公式
        # 三、分析讨论
        ('依据滚球法，保护半径和避雷针高度正相关', '[4]', 'end'), # 滚球法
        ('避雷针只能抵御直击雷，无法消除云层产生的感应静电', '[5]', 'end'), # 感应雷局限
        ('接地电阻必须控制在规定数值以内', '[4]', 'end'),         # 接地要求
        ('油罐、化工厂区域同样会利用尖端放电原理', '[5]', 'end'), # 工业拓展
        # 四、总结
        ('避雷针本质是利用尖端放电完成静电疏导', '[1]', 'end'),   # 核心原理
        ('防静电接地、屏蔽防护、静电消除器', '[5]', 'end'),       # 其他防护措施
    ]

    applied = 0
    for keyword, marker, mode in citations:
        for para in doc.paragraphs:
            if keyword in para.text:
                if mode == 'end':
                    # 在段落末尾run追加引用标记
                    if para.runs:
                        run = para.add_run(marker)
                        run.font.size = Pt(8)
                        run.font.superscript = True
                    applied += 1
                elif mode == 'after':
                    # 在关键词后插入标记（同段落内）
                    full = para.text
                    idx = full.find(keyword)
                    if idx >= 0 and para.runs:
                        # 用最简单方式：末尾加标记
                        run = para.add_run(' ' + marker)
                        run.font.size = Pt(8)
                        run.font.superscript = True
                    applied += 1
                break  # 每个关键词只匹配一次

    if applied:
        print(f'  已标注 {applied} 处正文引用')
    else:
        print('  [未找到可标注段落]')


def fix_latex_paragraphs(doc):
    """修复文档中的原始LaTeX代码"""
    fixed = 0
    for para in doc.paragraphs:
        text = para.text
        # P22: E\propto \frac{\sigma}{\varepsilon_0},\sigma=\frac{q}{S}
        if r'E\propto \frac{\sigma}{\varepsilon_0},\sigma=\frac{q}{S}' in text:
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = 'E = σ/ε0,   σ = q/S'
                para.runs[0].font.size = Pt(12)
                para.runs[0].font.bold = True
            fixed += 1
        # P23: 面电荷密度\sigma → 面电荷密度σ
        if r'\sigma' in text and para.runs:
            new_text = text.replace(r'\sigma', 'σ')
            for run in para.runs:
                run.text = ''
            para.runs[0].text = new_text
            fixed += 1
    if fixed:
        print(f'  已修复 {fixed} 处LaTeX残留')
    else:
        print('  [未找到LaTeX残留]')


def insert_image_after_paragraph(doc, keyword, img_path, width_inches, caption):
    """在包含keyword的段落后插入图片段落"""
    if not os.path.exists(img_path):
        print(f'  [文件不存在] {img_path}')
        return False

    target_para = None
    for para in doc.paragraphs:
        if keyword in para.text:
            target_para = para
            break

    if target_para is None:
        print(f'  [未找到关键词] "{keyword}"')
        return False

    # Insert XML elements
    img_p_elem = OxmlElement('w:p')
    target_para._element.addnext(img_p_elem)
    cap_p_elem = OxmlElement('w:p')
    img_p_elem.addnext(cap_p_elem)

    # Find the newly created paragraphs via doc.paragraphs iteration
    found_img_para = None
    found_cap_para = None
    for p in doc.paragraphs:
        if p._element is img_p_elem:
            found_img_para = p
        elif p._element is cap_p_elem:
            found_cap_para = p

    if found_img_para is None:
        print(f'  [无法定位新段落] {caption}')
        return False

    found_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    found_img_para.paragraph_format.space_before = Pt(14)
    found_img_para.paragraph_format.space_after = Pt(2)

    run = found_img_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    if found_cap_para:
        found_cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        found_cap_para.paragraph_format.space_after = Pt(14)
        cap_run = found_cap_para.add_run(caption)
        cap_run.font.size = Pt(9)

    print(f'  OK: {caption} ({width_inches}")')
    return True


def add_references(doc):
    """在文档末尾添加参考文献"""
    refs = [
        '[1] 赵凯华, 陈熙谋. 电磁学（第四版）[M]. 北京: 高等教育出版社, 2018.',
        '[2] 程守洙, 江之永. 普通物理学（第七版）上册[M]. 北京: 高等教育出版社, 2016.',
        '[3] 马文蔚, 周雨青. 物理学（第七版）上册[M]. 北京: 高等教育出版社, 2020.',
        '[4] GB 50057-2010, 建筑物防雷设计规范[S]. 北京: 中国建筑工业出版社, 2010.',
        '[5] GB/T 21714.1-2015, 雷电防护 第1部分: 总则[S]. 北京: 中国标准出版社, 2015.',
        '[6] 张大昌, 彭前程. 普通高中课程标准实验教科书 物理 选修3-1[M]. 北京: 人民教育出版社, 2010.',
    ]

    # Find the last non-empty paragraph (should be 总结 body text)
    last_para = None
    for p in doc.paragraphs:
        if p.text.strip():
            last_para = p

    if last_para is None:
        return

    # Create references heading
    head_elem = OxmlElement('w:p')
    last_para._element.addnext(head_elem)
    # Find it in doc
    for p in doc.paragraphs:
        if p._element is head_elem:
            p.style = doc.styles['Heading 1']
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run('参考文献')
            run.font.size = Pt(16)
            run.font.bold = True
            break

    # Add each reference
    prev_elem = head_elem
    for ref in refs:
        ref_elem = OxmlElement('w:p')
        prev_elem.addnext(ref_elem)
        prev_elem = ref_elem
        for p in doc.paragraphs:
            if p._element is ref_elem:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = Pt(-24)
                p.paragraph_format.left_indent = Pt(24)
                run = p.add_run(ref)
                run.font.size = Pt(10)
                break

    print('  已添加 6 条参考文献')


def main():
    print(f'读取: {SRC}')
    doc = Document(SRC)

    # 1. 修复LaTeX
    fix_latex_paragraphs(doc)

    # 2. 插入图片
    ok = 0
    for keyword, img_path, width, caption in IMAGES:
        if insert_image_after_paragraph(doc, keyword, img_path, width, caption):
            ok += 1
        else:
            print(f'  跳过: {caption}')

    # 3. 标注引用
    add_citations(doc)

    # 4. 添加参考文献
    add_references(doc)

    doc.save(OUT)
    print(f'\n完成: {ok}/{len(IMAGES)} 张图 -> {OUT}')
    print(f'文件大小: {os.path.getsize(OUT)//1024} KB')


if __name__ == '__main__':
    main()
